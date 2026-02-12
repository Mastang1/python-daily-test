#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
gen_word.py

Generate a Word document from ordered Markdown chapters.

Example:
    python gen_word.py --folder ./chapters --output ./output/design.docx
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# ============================================================
# Utilities
# ============================================================

def natural_sort_key(text: str):
    """
    Natural sorting: chapter2 < chapter10
    """
    return [
        int(part) if part.isdigit() else part.lower()
        for part in re.split(r"(\d+)", text)
    ]


def read_markdown_file(path: Path) -> str:
    """
    Read markdown file with BOM-safe UTF-8.
    """
    return path.read_text(encoding="utf-8-sig")


# ============================================================
# Markdown Loader
# ============================================================

class MarkdownLoader:
    """
    Load markdown files from a folder in natural order.
    """

    def __init__(self, folder: Path):
        self.folder = folder

    def load(self) -> List[Path]:
        if not self.folder.exists():
            raise FileNotFoundError(f"Folder not found: {self.folder}")

        md_files = [
            p for p in self.folder.iterdir()
            if p.is_file() and p.suffix.lower() == ".md"
        ]

        if not md_files:
            raise RuntimeError("No markdown files found in folder")

        return sorted(md_files, key=lambda p: natural_sort_key(p.name))


# ============================================================
# Word Document Builder
# ============================================================

class WordDocumentBuilder:
    """
    Build a Word document from Markdown text.
    """

    def __init__(self):
        self.document = Document()
        self._init_default_style()

    def _init_default_style(self):
        """
        Ensure Chinese-compatible default font.
        """
        style = self.document.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(11)

        # Ensure East Asia font is set (critical for Windows Word)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_chapter(self, chapter_title: str, markdown_text: str):
        """
        Add one markdown chapter into the Word document.
        """
        # Chapter heading
        self.document.add_heading(chapter_title, level=1)

        for raw_line in markdown_text.splitlines():
            line = raw_line.rstrip()

            if not line:
                self.document.add_paragraph("")
                continue

            # Markdown headings
            if line.startswith("#"):
                level = min(line.count("#"), 4)
                text = line.lstrip("#").strip()
                self.document.add_heading(text, level=level)
            else:
                p = self.document.add_paragraph(line)
                p.style = self.document.styles["Normal"]

    def save(self, output_path: Path):
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))


# ============================================================
# Workflow Orchestrator
# ============================================================

class GeneratorWorkflow:
    """
    Orchestrates the generation workflow.
    """

    def __init__(self, folder: Path, output: Path):
        self.folder = folder
        self.output = output

        self.loader = MarkdownLoader(folder)
        self.builder = WordDocumentBuilder()

    def run(self):
        md_files = self.loader.load()

        for md_file in md_files:
            content = read_markdown_file(md_file)
            chapter_title = md_file.stem
            self.builder.add_chapter(chapter_title, content)

        self.builder.save(self.output)


# ============================================================
# CLI
# ============================================================

def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate a Word document from Markdown chapters"
    )
    parser.add_argument(
        "--folder",
        required=True,
        help="Folder containing chapter markdown files"
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output Word file path (.docx)"
    )
    return parser


def main():
    # Make stdout safe on Windows
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    parser = build_arg_parser()
    args = parser.parse_args()

    folder = Path(args.folder).resolve()
    output = Path(args.output).resolve()

    workflow = GeneratorWorkflow(folder, output)
    workflow.run()

    print(f"[OK] Word document generated: {output}")


if __name__ == "__main__":
    main()
